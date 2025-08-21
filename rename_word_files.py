import os
from docx import Document
from docx.shared import Pt
import shutil

class RenameWordFiles:
    def __init__(self, type_id, facility, template_path):
        self.type_id = type_id
        self.template_file_path = None
        self.facility = facility
        self.word_file_path = None
        self.template_path = template_path
    
    def get_file(self):
        print(self.facility)
        print(self.type_id)
        word_file = None
        pdf_file = None
        
        self.template_file_path = os.path.join(self.template_path, self.facility, self.type_id)
        print(self.template_file_path)

        if os.path.exists(self.template_file_path):
            if self.type_id == 'Institutional':
                print(self.template_file_path)
                word_file = [os.path.join(self.template_file_path, f) 
                             for f in os.listdir(self.template_file_path)
                             if f.lower().endswith('.docx') and os.path.isfile(os.path.join(self.template_file_path, f))]
            
                print(word_file)

            elif self.type_id == 'Professional':
                pdf_file = [os.path.join(self.template_file_path, f) 
                            for f in os.listdir(self.template_file_path)
                            if f.lower().endswith('.pdf') and os.path.isfile(os.path.join(self.template_file_path, f))]
            
            if word_file and word_file[0] and os.path.exists(word_file[0]):
                self.word_file_path = word_file[0]
                print(self.word_file_path)

            if pdf_file and pdf_file[0] and os.path.exists(pdf_file[0]):
                self.word_file_path = pdf_file[0]
                print(self.word_file_path)


    def copy_word_file(self, dispute_id, dispute_no_path):
        rationale_file_path = ''
        print(self.word_file_path)
        if os.path.exists(self.word_file_path):
            if self.word_file_path.endswith('.pdf'):
                rationale_file_path = os.path.join(dispute_no_path, dispute_id + ' ACUITY AND RATIONALE.pdf')
            elif self.word_file_path.endswith('.docx'):
                rationale_file_path = os.path.join(dispute_no_path, dispute_id + ' ACUITY AND RATIONALE.docx')
            shutil.copy(self.word_file_path, rationale_file_path)
            

    def change_header(self, dispute_id, dispute_no_path):
        self.copy_word_file( dispute_id, dispute_no_path)
        if self.type_id == 'Institutional':
            doc = Document(self.word_file_path)
            header = header = doc.sections[0].header
            tables_in_header = header.tables
            print (dispute_no_path)
            print(dispute_id)
            if tables_in_header:
                 for table in tables_in_header:
                     first_row = table.rows[0]  
                     first_cell = first_row.cells[0]
                     new_paragraph = first_cell.add_paragraph()
                     new_paragraph.add_run(dispute_id)
                     for run in new_paragraph.runs:
                         run.font.name = 'Arial'  
                         run.font.size = Pt(11)
                     
            if os.path.exists(dispute_no_path):
                save_path = os.path.join(dispute_no_path, dispute_id + ' ACUITY AND RATIONALE.docx')
                doc.save(save_path)
                print(save_path)
                   